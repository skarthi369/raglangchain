import os
from typing import List, Tuple
from dotenv import load_dotenv
import pandas as pd
import json
from docx import Document
from PyPDF2 import PdfReader
from langchain_community.document_loaders import DirectoryLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.chains import ConversationalRetrievalChain
from langchain_groq import ChatGroq
from langchain.memory import ConversationBufferMemory
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document

# Load environment variables
load_dotenv()

class RAGChatbot:
    def __init__(self, data_dir: str = "data"):
        """
        Initialize the RAG Chatbot with data loading and processing.
        
        Args:
            data_dir (str): Directory containing the knowledge base files
        """
        self.data_dir = data_dir
        self.chain = None
        self.chat_history = []
        self._setup_rag_pipeline()
    
    def _process_feather_file(self, file_path: str) -> List[dict]:
        """Process a feather file and convert it to text format."""
        try:
            df = pd.read_feather(file_path)
            return self._dataframe_to_chunks(df, file_path)
        except Exception as e:
            print(f"Error processing feather file {file_path}: {str(e)}")
            return []
    
    def _process_csv_file(self, file_path: str) -> List[dict]:
        """Process a CSV file and convert it to text format."""
        try:
            df = pd.read_csv(file_path)
            return self._dataframe_to_chunks(df, file_path)
        except Exception as e:
            print(f"Error processing CSV file {file_path}: {str(e)}")
            return []
    
    def _process_excel_file(self, file_path: str) -> List[dict]:
        """Process an Excel file and convert it to text format."""
        try:
            df = pd.read_excel(file_path)
            return self._dataframe_to_chunks(df, file_path)
        except Exception as e:
            print(f"Error processing Excel file {file_path}: {str(e)}")
            return []
    
    def _process_json_file(self, file_path: str) -> List[dict]:
        """Process a JSON file and convert it to text format."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to text format
            if isinstance(data, list):
                # If JSON is a list of objects
                chunks = []
                for item in data:
                    if isinstance(item, dict):
                        chunk = " | ".join([f"{k}: {v}" for k, v in item.items()])
                        chunks.append({
                            "page_content": chunk,
                            "metadata": {
                                "source": file_path,
                                "type": "json",
                                "content_type": "list_item"
                            }
                        })
                return chunks
            elif isinstance(data, dict):
                # If JSON is a single object
                chunk = " | ".join([f"{k}: {v}" for k, v in data.items()])
                return [{
                    "page_content": chunk,
                    "metadata": {
                        "source": file_path,
                        "type": "json",
                        "content_type": "object"
                    }
                }]
            else:
                return [{
                    "page_content": str(data),
                    "metadata": {
                        "source": file_path,
                        "type": "json",
                        "content_type": "value"
                    }
                }]
        except Exception as e:
            print(f"Error processing JSON file {file_path}: {str(e)}")
            return []
    
    def _dataframe_to_chunks(self, df: pd.DataFrame, file_path: str) -> List[dict]:
        """Convert a DataFrame to text chunks."""
        chunks = []
        for _, row in df.iterrows():
            # Convert each row to a string representation
            row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
            chunks.append({
                "page_content": row_text,
                "metadata": {
                    "source": file_path,
                    "type": "dataframe",
                    "row_index": _,
                    "columns": list(df.columns)
                }
            })
        return chunks
    
    def _process_docx_file(self, file_path: str) -> List[dict]:
        """Process a Word document and convert it to text format."""
        try:
            doc = Document(file_path)
            chunks = []
            current_chunk = []
            current_length = 0
            section = "header"
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Check for section headers
                    if paragraph.style.name.startswith('Heading'):
                        if current_chunk:
                            chunks.append({
                                "page_content": "\n".join(current_chunk),
                                "metadata": {
                                    "source": file_path,
                                    "type": "docx",
                                    "section": section
                                }
                            })
                        current_chunk = [paragraph.text]
                        current_length = len(paragraph.text)
                        section = paragraph.text.lower()
                    else:
                        if current_length + len(paragraph.text) > 500:  # Reduced chunk size
                            chunks.append({
                                "page_content": "\n".join(current_chunk),
                                "metadata": {
                                    "source": file_path,
                                    "type": "docx",
                                    "section": section
                                }
                            })
                            current_chunk = [paragraph.text]
                            current_length = len(paragraph.text)
                        else:
                            current_chunk.append(paragraph.text)
                            current_length += len(paragraph.text)
            
            if current_chunk:
                chunks.append({
                    "page_content": "\n".join(current_chunk),
                    "metadata": {
                        "source": file_path,
                        "type": "docx",
                        "section": section
                    }
                })
            
            return chunks
        except Exception as e:
            print(f"Error processing Word document {file_path}: {str(e)}")
            return []
    
    def _process_pdf_file(self, file_path: str) -> List[dict]:
        """Process a PDF file and convert it to text format."""
        try:
            pdf = PdfReader(file_path)
            chunks = []
            current_chunk = []
            current_length = 0
            section = "header"
            
            for page in pdf.pages:
                text = page.extract_text()
                lines = text.split('\n')
                
                for line in lines:
                    if line.strip():
                        # Check for section headers (assuming headers are in all caps or have specific formatting)
                        if line.isupper() or line.strip().endswith(':'):
                            if current_chunk:
                                chunks.append({
                                    "page_content": "\n".join(current_chunk),
                                    "metadata": {
                                        "source": file_path,
                                        "type": "pdf",
                                        "section": section
                                    }
                                })
                            current_chunk = [line]
                            current_length = len(line)
                            section = line.lower()
                        else:
                            if current_length + len(line) > 500:  # Reduced chunk size
                                chunks.append({
                                    "page_content": "\n".join(current_chunk),
                                    "metadata": {
                                        "source": file_path,
                                        "type": "pdf",
                                        "section": section
                                    }
                                })
                                current_chunk = [line]
                                current_length = len(line)
                            else:
                                current_chunk.append(line)
                                current_length += len(line)
            
            if current_chunk:
                chunks.append({
                    "page_content": "\n".join(current_chunk),
                    "metadata": {
                        "source": file_path,
                        "type": "pdf",
                        "section": section
                    }
                })
            
            return chunks
        except Exception as e:
            print(f"Error processing PDF file {file_path}: {str(e)}")
            return []
    
    def _setup_rag_pipeline(self):
        """Set up the RAG pipeline with document loading, splitting, and embedding."""
        # Load and process documents
        documents = []
        
        # Process text files
        text_loader = DirectoryLoader(self.data_dir, glob="**/*.txt", loader_cls=TextLoader)
        documents.extend(text_loader.load())
        
        # Process other file types
        for file in os.listdir(self.data_dir):
            file_path = os.path.join(self.data_dir, file)
            if file.endswith('.feather'):
                chunks = self._process_feather_file(file_path)
                documents.extend([Document(page_content=chunk["page_content"], metadata=chunk["metadata"]) for chunk in chunks])
            elif file.endswith('.docx'):
                chunks = self._process_docx_file(file_path)
                documents.extend([Document(page_content=chunk["page_content"], metadata=chunk["metadata"]) for chunk in chunks])
            elif file.endswith('.pdf'):
                chunks = self._process_pdf_file(file_path)
                documents.extend([Document(page_content=chunk["page_content"], metadata=chunk["metadata"]) for chunk in chunks])
            elif file.endswith('.csv'):
                chunks = self._process_csv_file(file_path)
                documents.extend([Document(page_content=chunk["page_content"], metadata=chunk["metadata"]) for chunk in chunks])
            elif file.endswith('.json'):
                chunks = self._process_json_file(file_path)
                documents.extend([Document(page_content=chunk["page_content"], metadata=chunk["metadata"]) for chunk in chunks])
            elif file.endswith('.xlsx'):
                chunks = self._process_excel_file(file_path)
                documents.extend([Document(page_content=chunk["page_content"], metadata=chunk["metadata"]) for chunk in chunks])
        
        # Split documents into chunks with smaller size
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,  # Reduced chunk size
            chunk_overlap=100,  # Increased overlap
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )
        splits = text_splitter.split_documents(documents)
        
        # Create embeddings and vector store
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        vectorstore = Chroma.from_documents(
            documents=splits,
            embedding=embeddings
        )
        
        # Initialize memory for conversation history
        memory = ConversationBufferMemory(
            memory_key="chat_history",
            return_messages=True
        )
        
        # Create the conversational chain with Grok
        self.chain = ConversationalRetrievalChain.from_llm(
            llm=ChatGroq(
                model_name="meta-llama/llama-4-scout-17b-16e-instruct",
                temperature=0.7,
                groq_api_key=os.getenv("GROQ_API_KEY")
            ),
            retriever=vectorstore.as_retriever(
                search_kwargs={"k": 5}  # Increased number of retrieved documents
            ),
            memory=memory,
            combine_docs_chain_kwargs={
                "prompt": ChatPromptTemplate.from_template(
                    """You are a helpful AI assistant. Use the following context to answer the user's question.
                    If you don't know the answer, just say that you don't know. Don't try to make up an answer.
                    Pay special attention to contact information, personal details, and specific facts mentioned in the context.
                    
                    Context: {context}
                    
                    Chat History: {chat_history}
                    
                    Question: {question}
                    
                    Answer:"""
                )
            }
        )
    
    def get_response(self, query: str) -> str:
        """
        Get a response from the chatbot for a given query.
        
        Args:
            query (str): User's question or input
            
        Returns:
            str: Chatbot's response
        """
        if not self.chain:
            return "Error: RAG pipeline not initialized properly."
        
        result = self.chain({"question": query})
        self.chat_history.append((query, result["answer"]))
        return result["answer"]

def main():
    """Main function to run the chatbot in console mode."""
    print("Initializing RAG Chatbot...")
    chatbot = RAGChatbot()
    print("Chatbot initialized! Type 'quit' to exit.")
    
    while True:
        user_input = input("\nYou: ").strip()
        if user_input.lower() == 'quit':
            break
        
        response = chatbot.get_response(user_input)
        print(f"\nChatbot: {response}")

if __name__ == "__main__":
    main() 